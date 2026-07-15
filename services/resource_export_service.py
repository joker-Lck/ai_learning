"""学习资源文件导出服务
支持将不同类型的学习资源导出为相应的文件格式:
- 文档(document)和题库(quiz) → Word文档 (.docx)
- 思维导图(mindmap) → JPG图片 (.jpg)
- 视频(video)和动画(animation) → 视频脚本文件 (.txt)
- 代码案例(code_case)和阅读材料(reading) → Markdown文件 (.md)
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont

from core.logger import error, info, warning


class ResourceExportService:
    """学习资源文件导出服务"""

    def __init__(self):
        self.export_dir = "exports"
        os.makedirs(self.export_dir, exist_ok=True)
        info("资源导出服务初始化完成")

    def export_resource(self, resource: dict) -> dict | None:
        """
        根据资源类型导出为相应格式的文件

        Args:
            resource: 资源数据字典

        Returns:
            导出结果: {success, file_path, file_type, message}
        """
        try:
            resource_type = resource.get("type")
            content_data = resource.get("content_data", {})

            # 如果有 media_url（图片/视频/动画），直接返回该URL
            media_url = content_data.get("media_url")
            if media_url and media_url.startswith("/exports/"):
                file_path = media_url.replace("/exports/", "")
                full_path = os.path.join(self.export_dir, file_path)
                if os.path.exists(full_path):
                    return {
                        "success": True,
                        "file_path": full_path,
                        "file_type": file_path.split(".")[-1],
                        "filename": file_path,
                        "message": "导出成功"
                    }

            if resource_type in ["document", "quiz"]:
                return self._export_to_word(resource)
            elif resource_type == "mindmap":
                # 思维导图优先导出 SVG HTML
                if content_data.get("has_svg") and media_url:
                    return self._export_media(media_url)
                return self._export_mindmap_to_svg(resource)
            elif resource_type == "video":
                # 视频导出场景脚本 + 图片
                return self._export_video(resource)
            elif resource_type == "animation":
                # 图片导出
                if media_url:
                    return self._export_media(media_url)
                return self._export_to_markdown(resource)
            elif resource_type in ["code_case", "reading"]:
                return self._export_to_markdown(resource)
            else:
                warning(f"不支持的资源类型: {resource_type}")
                return {
                    "success": False,
                    "message": f"不支持的资源类型: {resource_type}"
                }

        except Exception as e:
            error(f"导出资源失败: {e!s}")
            return {
                "success": False,
                "message": f"导出失败: {e!s}"
            }

    def _export_media(self, media_url: str) -> dict:
        """导出媒体文件（图片/视频/动画）"""
        file_path = media_url.replace("/exports/", "")
        full_path = os.path.join(self.export_dir, file_path)
        if os.path.exists(full_path):
            return {
                "success": True,
                "file_path": full_path,
                "file_type": file_path.split(".")[-1],
                "filename": file_path,
                "message": "导出成功"
            }
        return {"success": False, "message": "文件不存在"}

    def _export_mindmap_to_svg(self, resource: dict) -> dict:
        """导出思维导图为 SVG"""
        content_data = resource.get("content_data", {})
        resource.get("title", "思维导图")

        # 如果有 SVG 数据，直接导出
        if content_data.get("svg_code"):
            filename = f"mindmap_{datetime.now().strftime('%Y%m%d_%H%M%S')}.svg"
            filepath = os.path.join(self.export_dir, filename)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content_data["svg_code"])
            return {
                "success": True,
                "file_path": filepath,
                "file_type": "svg",
                "filename": filename,
                "message": "导出成功"
            }

        # 否则导出为 JSON
        return self._export_to_markdown(resource)

    def _export_video(self, resource: dict) -> dict:
        """导出视频资源（脚本 + 图片）"""
        content_data = resource.get("content_data", {})
        resource.get("title", "教学视频")

        # 如果有 media_url，直接导出
        media_url = content_data.get("media_url")
        if media_url:
            return self._export_media(media_url)

        # 否则导出脚本为 Markdown
        return self._export_to_markdown(resource)

    def _export_to_word(self, resource: dict) -> dict:
        """导出文档或题库为Word格式"""
        try:
            doc = Document()

            # 添加标题
            title = doc.add_heading(resource.get("title", "学习资源"), level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加元信息
            meta_info = doc.add_paragraph()
            meta_info.add_run(f"学科: {resource.get('subject', 'N/A')}\n")
            meta_info.add_run(f"难度: {resource.get('difficulty_level', 'N/A')}\n")
            meta_info.add_run(f"生成时间: {resource.get('generated_at', 'N/A')}\n")

            content_data = resource.get("content_data", {})

            if resource["type"] == "document":
                # 导出文档内容
                sections = content_data.get("sections", [])
                for section in sections:
                    heading = section.get("heading", "")
                    content = section.get("content", "")

                    doc.add_heading(heading, level=2)
                    doc.add_paragraph(content)

                # 添加关键点
                key_points = content_data.get("key_points", [])
                if key_points:
                    doc.add_heading("关键要点", level=2)
                    for point in key_points:
                        doc.add_paragraph(point, style='List Bullet')

                # 添加参考资料
                references = content_data.get("references", [])
                if references:
                    doc.add_heading("参考资料", level=2)
                    for ref in references:
                        doc.add_paragraph(ref, style='List Number')

            elif resource["type"] == "quiz":
                # 导出题库内容
                questions = content_data.get("questions", [])
                for i, question in enumerate(questions, 1):
                    q_type = question.get("type", "unknown")
                    q_text = question.get("question", "")

                    # 题号和问题
                    doc.add_heading(f"第{i}题 [{self._get_question_type_name(q_type)}]", level=2)
                    doc.add_paragraph(q_text)

                    # 选项（如果是选择题）
                    options = question.get("options", [])
                    if options:
                        for option in options:
                            doc.add_paragraph(option)

                    # 答案
                    answer = question.get("answer", "")
                    doc.add_paragraph(f"答案: {answer}", style='Intense Quote')

                    # 解析
                    explanation = question.get("explanation", "")
                    if explanation:
                        doc.add_paragraph(f"解析: {explanation}")

                    # 知识点和难度
                    knowledge_point = question.get("knowledge_point", "")
                    difficulty = question.get("difficulty", "")
                    if knowledge_point or difficulty:
                        meta = doc.add_paragraph()
                        if knowledge_point:
                            meta.add_run(f"考察知识点: {knowledge_point}")
                        if difficulty:
                            meta.add_run(f" | 难度: {difficulty}")

            # 保存文件
            filename = f"{resource['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(self.export_dir, filename)
            doc.save(filepath)

            info(f"成功导出Word文档: {filepath}")
            return {
                "success": True,
                "file_path": filepath,
                "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "filename": filename,
                "message": "导出成功"
            }

        except Exception as e:
            error(f"导出Word文档失败: {e!s}")
            return {
                "success": False,
                "message": f"导出Word文档失败: {e!s}"
            }

    def _export_mindmap_to_jpg(self, resource: dict) -> dict:
        """导出思维导图为JPG图片"""
        try:
            content_data = resource.get("content_data", {})
            root = content_data.get("root", {})

            # 创建画布
            width, height = 1920, 1080
            img = Image.new('RGB', (width, height), color=(255, 255, 255))
            draw = ImageDraw.Draw(img)

            # 尝试使用系统字体
            try:
                font_title = ImageFont.truetype("simhei.ttf", 36)
                font_node = ImageFont.truetype("simsun.ttc", 24)
                font_small = ImageFont.truetype("simsun.ttc", 18)
            except:
                font_title = ImageFont.load_default()
                font_node = ImageFont.load_default()
                font_small = ImageFont.load_default()

            # 绘制标题
            title = resource.get("title", "思维导图")
            draw.text((width//2 - 200, 50), title, fill=(0, 0, 0), font=font_title)

            # 递归绘制节点
            self._draw_tree_node(draw, root, width//2, 150, width-100, font_node, font_small, level=0)

            # 保存文件
            filename = f"mindmap_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg"
            filepath = os.path.join(self.export_dir, filename)
            img.save(filepath, 'JPEG', quality=95)

            info(f"成功导出思维导图JPG: {filepath}")
            return {
                "success": True,
                "file_path": filepath,
                "file_type": "image/jpeg",
                "filename": filename,
                "message": "导出成功"
            }

        except Exception as e:
            error(f"导出思维导图JPG失败: {e!s}")
            return {
                "success": False,
                "message": f"导出思维导图JPG失败: {e!s}"
            }

    def _draw_tree_node(self, draw, node, x, y, max_width, font_node, font_small, level=0):
        """递归绘制树形节点"""
        if not node:
            return

        name = node.get("name", "")
        children = node.get("children", [])

        # 绘制当前节点
        node_color = self._get_level_color(level)
        text_bbox = draw.textbbox((0, 0), name, font=font_node)
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]

        # 绘制圆角矩形背景
        padding = 15
        rect_x1 = x - text_width//2 - padding
        rect_y1 = y - text_height//2 - padding
        rect_x2 = x + text_width//2 + padding
        rect_y2 = y + text_height//2 + padding

        draw.rounded_rectangle([rect_x1, rect_y1, rect_x2, rect_y2],
                              radius=10, fill=node_color, outline=(0, 0, 0))

        # 绘制文本
        text_x = x - text_width//2
        text_y = y - text_height//2
        draw.text((text_x, text_y), name, fill=(255, 255, 255), font=font_node)

        # 递归绘制子节点
        if children:
            child_spacing = max_width // (len(children) + 1)
            for i, child in enumerate(children):
                child_x = child_spacing * (i + 1)
                child_y = y + 120

                # 绘制连接线
                draw.line([(x, rect_y2), (child_x, child_y - 40)],
                         fill=(100, 100, 100), width=2)

                # 递归绘制子节点
                self._draw_tree_node(draw, child, child_x, child_y,
                                   child_spacing, font_node, font_small, level + 1)

    def _get_level_color(self, level):
        """根据层级获取颜色"""
        colors = [
            (70, 130, 180),   # 钢蓝色 - 根节点
            (34, 139, 34),    # 森林绿 - 第一层
            (255, 140, 0),    # 深橙色 - 第二层
            (138, 43, 226),   # 紫罗兰 - 第三层
            (220, 20, 60),    # 深红色 - 第四层
        ]
        return colors[min(level, len(colors) - 1)]

    def _export_script_to_txt(self, resource: dict) -> dict:
        """导出视频/动画脚本为TXT文件"""
        try:
            content_data = resource.get("content_data", {})

            lines = []
            lines.append(f"{'='*60}")
            lines.append(f"标题: {resource.get('title', '脚本')}")
            lines.append(f"学科: {resource.get('subject', 'N/A')}")
            lines.append(f"难度: {resource.get('difficulty_level', 'N/A')}")
            lines.append(f"时长: {content_data.get('duration_minutes', 'N/A')} 分钟")
            lines.append(f"生成时间: {resource.get('generated_at', 'N/A')}")
            lines.append(f"{'='*60}\n")

            if resource["type"] == "video":
                # 视频脚本
                scenes = content_data.get("scenes", [])
                lines.append("【场景列表】\n")

                for scene in scenes:
                    scene_id = scene.get("scene_id", 0)
                    duration = scene.get("duration_seconds", 0)
                    visual = scene.get("visual_description", "")
                    narration = scene.get("narration", "")
                    effects = scene.get("animation_effects", [])

                    lines.append(f"场景 {scene_id} (时长: {duration}秒)")
                    lines.append(f"画面: {visual}")
                    lines.append(f"旁白: {narration}")
                    if effects:
                        lines.append(f"特效: {', '.join(effects)}")
                    lines.append("")

                # 关键画面
                key_visuals = content_data.get("key_visuals", [])
                if key_visuals:
                    lines.append("\n【关键画面】")
                    for visual in key_visuals:
                        lines.append(f"- {visual}")

            elif resource["type"] == "animation":
                # 动画脚本
                frames = content_data.get("frames", [])
                lines.append("【动画帧列表】\n")

                for frame in frames:
                    frame_id = frame.get("frame_id", 0)
                    timestamp = frame.get("timestamp", "")
                    description = frame.get("description", "")
                    action = frame.get("action", "")
                    transition = frame.get("transition", "")

                    lines.append(f"帧 {frame_id} (时间: {timestamp})")
                    lines.append(f"描述: {description}")
                    lines.append(f"动作: {action}")
                    if transition:
                        lines.append(f"转场: {transition}")
                    lines.append("")

                # 解说词
                narration = content_data.get("narration_script", "")
                if narration:
                    lines.append("\n【完整解说词】")
                    lines.append(narration)

                # 视觉风格
                visual_style = content_data.get("visual_style", "")
                if visual_style:
                    lines.append(f"\n【视觉风格】{visual_style}")

            # 保存文件
            filename = f"{resource['type']}_script_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            filepath = os.path.join(self.export_dir, filename)

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write('\n'.join(lines))

            info(f"成功导出脚本文件: {filepath}")
            return {
                "success": True,
                "file_path": filepath,
                "file_type": "text/plain",
                "filename": filename,
                "message": "导出成功"
            }

        except Exception as e:
            error(f"导出脚本文件失败: {e!s}")
            return {
                "success": False,
                "message": f"导出脚本文件失败: {e!s}"
            }

    def _export_to_markdown(self, resource: dict) -> dict:
        """导出代码案例或阅读材料为Markdown文件"""
        try:
            content_data = resource.get("content_data", {})

            lines = []
            lines.append(f"# {resource.get('title', '学习资源')}\n")
            lines.append(f"**学科**: {resource.get('subject', 'N/A')}  ")
            lines.append(f"**难度**: {resource.get('difficulty_level', 'N/A')}  ")
            lines.append(f"**生成时间**: {resource.get('generated_at', 'N/A')}  \n")
            lines.append("---\n")

            if resource["type"] == "code_case":
                # 代码案例
                description = content_data.get("description", "")
                if description:
                    lines.append(f"## 案例说明\n\n{description}\n")

                requirements = content_data.get("requirements", [])
                if requirements:
                    lines.append("## 需求说明\n")
                    for req in requirements:
                        lines.append(f"- {req}")
                    lines.append("")

                implementation_steps = content_data.get("implementation_steps", [])
                if implementation_steps:
                    lines.append("## 实现步骤\n")
                    for i, step in enumerate(implementation_steps, 1):
                        lines.append(f"{i}. {step}")
                    lines.append("")

                # 代码部分
                code_info = content_data.get("code", {})
                if code_info:
                    language = code_info.get("language", "text")
                    filename = code_info.get("filename", "code.txt")
                    source_code = code_info.get("source_code", "")

                    lines.append(f"## 代码实现 (`{filename}`)\n")
                    lines.append(f"```{language}")
                    lines.append(source_code)
                    lines.append("```\n")

                expected_output = content_data.get("expected_output", "")
                if expected_output:
                    lines.append("## 预期输出\n")
                    lines.append("```\n" + expected_output + "\n```\n")

                exercises = content_data.get("exercises", [])
                if exercises:
                    lines.append("## 扩展练习\n")
                    for exercise in exercises:
                        lines.append(f"- {exercise}")
                    lines.append("")

            elif resource["type"] == "reading":
                # 阅读材料
                content = content_data.get("content", "")
                if content:
                    lines.append(f"## 正文内容\n\n{content}\n")

                case_studies = content_data.get("case_studies", [])
                if case_studies:
                    lines.append("## 案例分析\n")
                    for case in case_studies:
                        lines.append(f"- {case}")
                    lines.append("")

                further_reading = content_data.get("further_reading", [])
                if further_reading:
                    lines.append("## 延伸阅读\n")
                    for item in further_reading:
                        title = item.get("title", "")
                        url = item.get("url", "")
                        if url:
                            lines.append(f"- [{title}]({url})")
                        else:
                            lines.append(f"- {title}")
                    lines.append("")

                references = content_data.get("references", [])
                if references:
                    lines.append("## 参考文献\n")
                    for ref in references:
                        lines.append(f"1. {ref}")
                    lines.append("")

            # 保存文件
            filename = f"{resource['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
            filepath = os.path.join(self.export_dir, filename)

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write('\n'.join(lines))

            info(f"成功导出Markdown文件: {filepath}")
            return {
                "success": True,
                "file_path": filepath,
                "file_type": "text/markdown",
                "filename": filename,
                "message": "导出成功"
            }

        except Exception as e:
            error(f"导出Markdown文件失败: {e!s}")
            return {
                "success": False,
                "message": f"导出Markdown文件失败: {e!s}"
            }

    def _get_question_type_name(self, q_type: str) -> str:
        """获取题目类型的中文名称"""
        type_map = {
            "multiple_choice": "选择题",
            "fill_in_blank": "填空题",
            "essay": "解答题",
            "true_false": "判断题",
            "short_answer": "简答题"
        }
        return type_map.get(q_type, q_type)


# 全局实例
resource_export_service = ResourceExportService()
